#!/usr/bin/env python3
"""Generate a branded Nhancio application-development quotation from JSON data."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from generate_offer_letter import (
    BASE_DIR,
    CONTENT_WIDTH_DXA,
    DARK_PURPLE,
    INK,
    LIGHT_FILL,
    MUTED,
    OfferValidationError,
    PURPLE,
    WHITE,
    add_bullet,
    add_label_value_paragraph,
    add_picture_alt,
    add_term,
    configure_page_and_furniture,
    configure_styles,
    create_bullet_numbering,
    display_date,
    load_json,
    parse_iso_date,
    required,
    safe_filename,
    set_cell_border,
    set_cell_shading,
    set_run_font,
    set_table_geometry,
)


def indian_group(amount: int | float) -> str:
    value = int(round(amount))
    sign = "-" if value < 0 else ""
    digits = str(abs(value))
    if len(digits) <= 3:
        grouped = digits
    else:
        last_three = digits[-3:]
        rest = digits[:-3]
        parts: list[str] = []
        while rest:
            parts.append(rest[-2:])
            rest = rest[:-2]
        grouped = ",".join(reversed(parts)) + "," + last_three
    return f"{sign}{grouped}"


def inr_amount(amount: int | float) -> str:
    return f"INR {indian_group(amount)}"


def validate_quotation(profile: dict[str, Any], quote: dict[str, Any]) -> None:
    for field in (
        "display_name",
        "legal_name",
        "website",
        "address_lines",
        "logo_image",
        "signatory.name",
        "signatory.title",
        "signatory.signature_image",
    ):
        required(profile, field)

    for field in (
        "reference_number",
        "issue_date",
        "valid_until",
        "client.name",
        "project.name",
        "project.subtitle",
        "project.summary",
        "project.work_mode",
        "project.location",
        "project.engagement_type",
        "project.proposed_start",
        "project.build_duration",
        "commercials.one_time_fee",
        "commercials.monthly_fee",
        "commercials.gst_note",
        "commercials.one_time_allocation",
        "commercials.payment_schedule",
        "scope.modules",
        "scope.managed_services",
        "scope.exclusions",
        "scope.assumptions",
        "terms.validity",
        "terms.notice_period",
        "terms.governing_law",
    ):
        required(quote, field)

    issue = parse_iso_date(quote["issue_date"], "issue_date")
    valid_until = parse_iso_date(quote["valid_until"], "valid_until")
    parse_iso_date(quote["project"]["proposed_start"], "project.proposed_start")
    if valid_until < issue:
        raise OfferValidationError("valid_until cannot be before issue_date")

    one_time = quote["commercials"]["one_time_fee"]
    monthly = quote["commercials"]["monthly_fee"]
    if not isinstance(one_time, (int, float)) or one_time <= 0:
        raise OfferValidationError("commercials.one_time_fee must be a positive number")
    if not isinstance(monthly, (int, float)) or monthly <= 0:
        raise OfferValidationError("commercials.monthly_fee must be a positive number")

    allocation = quote["commercials"]["one_time_allocation"]
    if not isinstance(allocation, list) or not allocation:
        raise OfferValidationError("commercials.one_time_allocation must be a non-empty list")
    allocated = 0
    for row in allocation:
        if not isinstance(row, dict) or not row.get("item") or "amount" not in row:
            raise OfferValidationError("Each allocation row needs item and amount")
        allocated += int(row["amount"])
    if allocated != int(one_time):
        raise OfferValidationError(
            f"Allocation total {allocated} does not match one-time fee {int(one_time)}"
        )

    for key in ("modules", "managed_services", "exclusions", "assumptions"):
        items = quote["scope"][key]
        if not isinstance(items, list) or not all(isinstance(item, str) and item.strip() for item in items):
            raise OfferValidationError(f"scope.{key} must be a non-empty list of text items")

    logo = (BASE_DIR / profile["logo_image"]).resolve()
    signature = (BASE_DIR / profile["signatory"]["signature_image"]).resolve()
    if not logo.is_file():
        raise OfferValidationError(f"Missing asset: {logo}")
    if not signature.is_file():
        raise OfferValidationError(f"Missing asset: {signature}")


def loosen_quotation_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    heading_tokens = {
        "Heading 1": (14, 12, 4),
        "Heading 2": (11.5, 9, 3),
        "Heading 3": (10.5, 7, 3),
    }
    for style_name, (size, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    emphasize_last_row: bool = False,
) -> None:
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_geometry(table, widths_dxa)
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "5B2DCB")
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)
        if idx > 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run(header)
        set_run_font(run, size=9, color=WHITE, bold=True)
    for row_idx, values in enumerate(rows):
        for col_idx, value in enumerate(values):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if emphasize_last_row and row_idx == len(rows) - 1:
                set_cell_shading(cell, LIGHT_FILL)
            elif row_idx % 2 == 1:
                set_cell_shading(cell, "FBFAFE")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            if col_idx > 0 and col_idx == len(values) - 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            bold = emphasize_last_row and row_idx == len(rows) - 1
            color = DARK_PURPLE if bold else INK
            run = paragraph.add_run(value)
            set_run_font(run, size=9.5, color=color, bold=bold)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_key_commercials_table(document: Document, quote: dict[str, Any]) -> None:
    commercials = quote["commercials"]
    project = quote["project"]
    rows = [
        ("Project", f"{project['name']} — {project['subtitle']}"),
        ("Client", quote["client"]["name"]),
        ("Delivery location", project["location"]),
        ("Engagement", project["engagement_type"]),
        ("Work arrangement", project["work_mode"]),
        ("Proposed start", display_date(project["proposed_start"], "project.proposed_start")),
        ("Build window", project["build_duration"]),
        ("One-time build fee", inr_amount(commercials["one_time_fee"])),
        ("Monthly managed services", f"{inr_amount(commercials['monthly_fee'])} per month"),
        ("Quotation valid until", display_date(quote["valid_until"], "valid_until")),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [3120, CONTENT_WIDTH_DXA - 3120])
    table.style = "Table Grid"
    for row_idx, (label, value) in enumerate(rows):
        left, right = table.rows[row_idx].cells
        set_cell_shading(left, LIGHT_FILL)
        if row_idx % 2 == 1:
            set_cell_shading(right, "FBFAFE")
        for cell, text_value, bold, color in (
            (left, label, True, DARK_PURPLE),
            (right, value, False, INK),
        ):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(text_value)
            set_run_font(run, size=9.5, color=color, bold=bold)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_client_acceptance(document: Document, client_name: str) -> None:
    table = document.add_table(rows=3, cols=2)
    set_table_geometry(table, [CONTENT_WIDTH_DXA - 3600, 3600])
    labels = (
        (f"Authorized signatory, {client_name}", "Date"),
        ("Name / title", "Place"),
    )
    for col, (top, bottom) in enumerate(zip(labels[0], labels[1], strict=True)):
        line_cell = table.rows[0].cells[col]
        line_p = line_cell.paragraphs[0]
        line_p.paragraph_format.space_before = Pt(18)
        line_p.paragraph_format.space_after = Pt(2)
        line_p.add_run(" ")
        set_cell_border(line_cell, bottom={"val": "single", "sz": 8, "space": 0, "color": "6B6E78"})
        name_cell = table.rows[1].cells[col]
        name_p = name_cell.paragraphs[0]
        name_p.paragraph_format.space_before = Pt(2)
        name_p.paragraph_format.space_after = Pt(0)
        run = name_p.add_run(top)
        set_run_font(run, size=9, color=MUTED)
        extra = table.rows[2].cells[col]
        extra_p = extra.paragraphs[0]
        extra_p.paragraph_format.space_before = Pt(10)
        extra_p.paragraph_format.space_after = Pt(0)
        extra_run = extra_p.add_run(bottom)
        set_run_font(extra_run, size=9, color=MUTED)
        set_cell_border(name_cell, top={"val": "nil"}, bottom={"val": "nil"})
        set_cell_border(extra, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})


def build_quotation_document(profile: dict[str, Any], quote: dict[str, Any]) -> Document:
    document = Document()
    configure_styles(document)
    loosen_quotation_styles(document)
    configure_page_and_furniture(document, profile, document_label="Application Development Quotation")
    bullet_num_id = create_bullet_numbering(document)

    project = quote["project"]
    commercials = quote["commercials"]
    client = quote["client"]
    one_time = int(commercials["one_time_fee"])
    monthly = int(commercials["monthly_fee"])

    document.core_properties.title = f"Quotation - {project['name']}"
    document.core_properties.subject = project["subtitle"]
    document.core_properties.author = profile["legal_name"]
    document.core_properties.company = profile["legal_name"]

    if quote.get("document_status"):
        status = document.add_paragraph()
        status.alignment = WD_ALIGN_PARAGRAPH.CENTER
        status.paragraph_format.space_after = Pt(4)
        run = status.add_run(quote["document_status"])
        set_run_font(run, size=8, color=WHITE, bold=True)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "6B2FD0")
        status._p.get_or_add_pPr().append(shading)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("APPLICATION DEVELOPMENT QUOTATION")
    set_run_font(title_run, size=16, color=DARK_PURPLE, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    sub_run = subtitle.add_run(f"{project['name']}  ·  {project['subtitle']}")
    set_run_font(sub_run, size=11, color=PURPLE, bold=True)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.paragraph_format.space_after = Pt(1)
    set_run_font(meta.add_run(display_date(quote["issue_date"], "issue_date")), size=10, color=INK)
    ref = document.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ref.paragraph_format.space_after = Pt(8)
    set_run_font(ref.add_run(f"Reference: {quote['reference_number']}"), size=9.5, color=MUTED)

    add_label_value_paragraph(document, "To", client["name"])
    for line in client.get("address_lines") or []:
        p = document.add_paragraph(line)
        p.paragraph_format.space_after = Pt(1)
    if client.get("contact_name"):
        add_label_value_paragraph(document, "Attention", client["contact_name"])

    subject = document.add_paragraph()
    subject.paragraph_format.space_before = Pt(6)
    subject.paragraph_format.space_after = Pt(6)
    set_run_font(
        subject.add_run(
            f"Subject: Quotation for {project['name']} — {project['subtitle']}"
        ),
        size=10.5,
        color=DARK_PURPLE,
        bold=True,
    )

    greeting_name = client.get("contact_name") or "Team"
    salutation = document.add_paragraph(f"Dear {greeting_name},")
    salutation.paragraph_format.space_after = Pt(6)
    intro = document.add_paragraph()
    intro.add_run("Nhancio is pleased to quote for the design, development, deployment, and ongoing operation of ")
    bold = intro.add_run(project["name"])
    bold.bold = True
    intro.add_run(
        f", a Hyderabad worker application modelled on Pronto-style shift booking and SnapIt-style nearby-job discovery. "
        f"{project['summary']}"
    )

    document.add_paragraph("1. Commercial snapshot", style="Heading 1")
    add_key_commercials_table(document, quote)
    note = document.add_paragraph(commercials["gst_note"])
    note.paragraph_format.space_after = Pt(6)
    set_run_font(note.runs[0], size=9.5, color=MUTED, italic=True)

    document.add_paragraph("2. Project understanding", style="Heading 1")
    add_term(
        document,
        "What Worker X is",
        "Worker X is a worker-facing marketplace for Hyderabad. It copies the practical worker loop of Pronto (see available shift, claim it, check in, get paid) and the nearby-job loop of SnapIt (browse local work, apply quickly, track status). The first release is for workers in Hyderabad, not a national multi-city platform and not an employer-suite rebuild.",
    )
    add_term(
        document,
        "Why this quotation is structured this way",
        f"The one-time fee of {inr_amount(one_time)} covers the build, server setup, deployment, localization, and launch support. The monthly fee of {inr_amount(monthly)} is a managed-services retainer for bug management and content updates after go-live. New feature work, extra products, or paid media campaigns are outside this quotation.",
    )

    document.add_paragraph("3. Scope of work — first release", style="Heading 1")
    mapping_intro = document.add_paragraph(
        "The first release maps Pronto and SnapIt worker capabilities onto a single Hyderabad product. The table below is the agreed reference, not a licence of those brands."
    )
    mapping_intro.paragraph_format.space_after = Pt(6)
    add_header_table(
        document,
        ["Capability", "Pronto-style", "SnapIt-style", "Worker X (Hyderabad)"],
        [
            [
                "Find work",
                "Open shifts by time and site",
                "Nearby jobs on a map/list",
                "Hyderabad zone feed of shifts and nearby jobs",
            ],
            [
                "Take work",
                "Claim / confirm a shift",
                "Quick apply from the job card",
                "Claim shift or apply to a local job in-app",
            ],
            [
                "Do work",
                "Geo check-in / check-out",
                "Status updates on the assignment",
                "Check-in, check-out, and assignment status",
            ],
            [
                "Get paid",
                "Shift earnings ledger",
                "Job payout status",
                "Earnings view with UPI-oriented payout states",
            ],
            [
                "Stay informed",
                "Shift reminders",
                "New-job alerts",
                "Push alerts in English and Telugu copy fields",
            ],
        ],
        [2200, 2500, 2500, CONTENT_WIDTH_DXA - 7200],
    )

    document.add_paragraph("Included modules", style="Heading 2")
    for item in quote["scope"]["modules"]:
        add_bullet(document, item, bullet_num_id)

    document.add_paragraph("4. One-time fee allocation", style="Heading 1")
    alloc_intro = document.add_paragraph(
        f"The one-time professional fee is {inr_amount(one_time)}. It is allocated as follows so server, deployment, development, and related launch work are visible as separate lines. This is an internal allocation of a single fixed fee, not a menu of optional add-ons."
    )
    alloc_intro.paragraph_format.space_after = Pt(6)

    allocation_rows: list[list[str]] = []
    for row in commercials["one_time_allocation"]:
        amount = int(row["amount"])
        share = f"{round(amount / one_time * 100)}%"
        allocation_rows.append(
            [row["item"], row.get("notes") or "", inr_amount(amount), share]
        )
    allocation_rows.append(["Total one-time fee", "Fixed-price Worker X first release", inr_amount(one_time), "100%"])
    add_header_table(
        document,
        ["Line item", "What it covers", "Amount", "Share"],
        allocation_rows,
        [2400, CONTENT_WIDTH_DXA - 2400 - 1800 - 1100, 1800, 1100],
        emphasize_last_row=True,
    )

    add_term(
        document,
        "How to read the allocation",
        "Development is the largest line because Worker X has to reproduce the worker loops of Pronto and SnapIt in one Hyderabad app. Server cost is setup and first-period environment work, not an open-ended cloud bill. Deployment covers production release and store listing support. Product/UX/localization and QA are the remaining launch costs. Recurring hosting overages, SMS, and gateway fees remain extra as listed in exclusions.",
    )

    document.add_paragraph("5. Monthly managed services", style="Heading 1")
    monthly_intro = document.add_paragraph(
        f"From the month of go-live, Nhancio will provide managed services at {inr_amount(monthly)} per month, billed in advance. This retainer is for keeping Worker X stable and current, not for building a second product."
    )
    monthly_intro.paragraph_format.space_after = Pt(6)
    add_header_table(
        document,
        ["Service", "Included each month"],
        [
            ["Bug management", "Triage, reproduce, and fix in-scope defects in the live Worker X application. Priority incidents are handled first during business hours."],
            ["Content updates", "Job categories, banners, FAQs, Hyderabad city copy, and approved store-listing text, published after a short staging check."],
            ["Minor patches", "Small production publishes of approved fixes and copy. Not new modules, redesigns, or extra city launches."],
            ["Reporting", "A monthly health note covering incidents closed, content changed, and items still open."],
        ],
        [2400, CONTENT_WIDTH_DXA - 2400],
    )
    for item in quote["scope"]["managed_services"]:
        add_bullet(document, item, bullet_num_id)
    add_term(
        document,
        "What the monthly fee does not include",
        "New screens, extra cities, employer-console expansion, paid user-acquisition, or statutory payroll processing. Those are quoted separately if required. Either party may end managed services by giving "
        f"{quote['terms']['notice_period']} written notice after go-live.",
    )

    document.add_paragraph("6. Payment schedule", style="Heading 1")
    pay_rows = []
    for row in commercials["payment_schedule"]:
        pay_rows.append([row["milestone"], row.get("notes") or "", inr_amount(int(row["amount"]))])
    add_header_table(
        document,
        ["When", "Purpose", "Amount"],
        pay_rows,
        [3800, CONTENT_WIDTH_DXA - 3800 - 1800, 1800],
    )
    add_term(
        document,
        "Invoicing",
        commercials["gst_note"]
        + " The one-time fee is due as stated above. Monthly managed services are invoiced in advance for the coming month. Work on a milestone begins after the related payment is received, unless Nhancio agrees otherwise in writing.",
    )

    document.add_paragraph("7. Delivery approach", style="Heading 1")
    add_header_table(
        document,
        ["Stage", "Nhancio delivers", "Client provides"],
        [
            [
                "Kickoff",
                "Feature map from Pronto/SnapIt to Worker X, environment plan, and week-by-week build list",
                "Brand assets, Telugu/English copy, zone list, and store/payment accounts",
            ],
            [
                "Build",
                "Worker app modules in the agreed first-release scope, with staging access",
                "Review comments within two working days of each staging drop",
            ],
            [
                "UAT",
                "Test notes, defect closure for in-scope items, and a go-live checklist",
                "Named reviewer, sample jobs/shifts, and written UAT sign-off",
            ],
            [
                "Go-live",
                "Production deployment, store listing support, and handover of credentials in the client name",
                "Store approval, final content, and the remaining one-time payment",
            ],
            [
                "Operate",
                "Bug management and content updates under the monthly retainer",
                "A single content owner and a written change request for each update",
            ],
        ],
        [1600, (CONTENT_WIDTH_DXA - 1600) // 2, CONTENT_WIDTH_DXA - 1600 - (CONTENT_WIDTH_DXA - 1600) // 2],
    )
    add_term(
        document,
        "Build window",
        project["build_duration"]
        + " Dates move if copy, accounts, or approvals are late. Nhancio will flag slippage as soon as it is visible.",
    )

    document.add_paragraph("8. Assumptions and exclusions", style="Heading 1")
    document.add_paragraph("Assumptions", style="Heading 2")
    for item in quote["scope"]["assumptions"]:
        add_bullet(document, item, bullet_num_id)
    document.add_paragraph("Not included", style="Heading 2")
    for item in quote["scope"]["exclusions"]:
        add_bullet(document, item, bullet_num_id)

    document.add_paragraph("9. Commercial and legal terms", style="Heading 1")
    add_term(
        document,
        "Validity",
        quote["terms"]["validity"]
        + f" Unless accepted by {display_date(quote['valid_until'], 'valid_until')}, Nhancio may revise fees or withdraw this quotation.",
    )
    add_term(
        document,
        "Intellectual property",
        "On receipt of the one-time fee in full, the client receives a licence to use the delivered Worker X application for its Hyderabad operations. Nhancio retains pre-existing tools, libraries, and internal accelerators. Third-party services (stores, maps, SMS, UPI) remain subject to their own terms.",
    )
    add_term(
        document,
        "Confidentiality",
        "Each party will keep the other party's non-public business, product, and user information confidential and use it only to perform this engagement. This quotation itself is confidential to the named client and Nhancio.",
    )
    add_term(
        document,
        "Change control",
        "Material additions to scope — extra modules, extra cities, or employer-console work — require a written change note and a revised fee. Bug management and content updates inside the monthly retainer do not.",
    )
    add_term(
        document,
        "Limitation and governing law",
        "Nhancio's aggregate liability under this quotation is limited to the fees paid for the relevant one-time or monthly period, except where applicable law does not allow that limit. This quotation is governed by the laws applicable in "
        f"{quote['terms']['governing_law']}, and the competent courts there will have jurisdiction, subject to any mandatory law that applies.",
    )

    document.add_paragraph("10. Authorization", style="Heading 1")
    close = document.add_paragraph(
        "We are ready to start Worker X as a focused Hyderabad copy of the Pronto and SnapIt worker experience, on the fees and allocation set out above. Please sign below to accept this quotation."
    )
    close.paragraph_format.space_after = Pt(8)

    signatory = profile["signatory"]
    auth = document.add_paragraph()
    auth.paragraph_format.space_after = Pt(2)
    set_run_font(auth.add_run(f"For and on behalf of {profile['legal_name']}"), size=10.5, color=DARK_PURPLE, bold=True)
    sig_p = document.add_paragraph()
    sig_p.paragraph_format.space_after = Pt(2)
    shape = sig_p.add_run().add_picture(str(BASE_DIR / signatory["signature_image"]), width=Inches(1.25))
    add_picture_alt(shape, "Authorized signature", f"Signature of {signatory['name']}")
    name_p = document.add_paragraph()
    name_p.paragraph_format.space_after = Pt(1)
    set_run_font(name_p.add_run(signatory["name"]), size=10.5, color=INK, bold=True)
    add_label_value_paragraph(document, "Title", signatory["title"])
    add_label_value_paragraph(document, "Date", display_date(quote["issue_date"], "issue_date"))
    add_label_value_paragraph(document, "Place", "Hyderabad")

    accept = document.add_paragraph("Client acceptance", style="Heading 2")
    accept.paragraph_format.space_before = Pt(14)
    accept_body = document.add_paragraph(
        f"By signing, {client['name']} accepts this quotation, the {inr_amount(one_time)} one-time fee allocation, and the {inr_amount(monthly)} monthly managed-services fee for bug management and content updates."
    )
    accept_body.paragraph_format.space_after = Pt(4)
    add_client_acceptance(document, client["name"])
    return document


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--offer", "--quote", dest="quote", type=Path, default=BASE_DIR / "sample_quotation.json")
    parser.add_argument("--profile", type=Path, default=BASE_DIR / "company_profile.json")
    parser.add_argument("--output", type=Path)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    try:
        profile = load_json(args.profile.resolve())
        quote = load_json(args.quote.resolve())
        validate_quotation(profile, quote)
        output = (
            args.output
            or BASE_DIR / "generated" / f"{safe_filename(quote['project']['name'])}_Application_Development_Quotation.docx"
        ).resolve()
        if output.suffix.lower() != ".docx":
            raise OfferValidationError("Output filename must end in .docx")
        output.parent.mkdir(parents=True, exist_ok=True)
        build_quotation_document(profile, quote).save(output)
        print(output)
        return 0
    except OfferValidationError as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
