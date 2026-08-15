#!/usr/bin/env python3
"""Generate a branded Nhancio employment offer letter from JSON data."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from generate_offer_letter import (
    BASE_DIR,
    CONTENT_WIDTH_DXA,
    DARK_PURPLE,
    INK,
    MUTED,
    PURPLE,
    WHITE,
    OfferValidationError,
    add_bullet,
    add_key_terms_table,
    add_label_value_paragraph,
    add_picture_alt,
    add_term,
    build_document,
    configure_page_and_furniture,
    configure_styles,
    create_bullet_numbering,
    display_date,
    load_json,
    safe_filename,
    set_run_font,
    set_table_geometry,
    validate,
)


def validate_employment(profile: dict[str, Any], offer: dict[str, Any]) -> None:
    required_fields = [
        "reference_number",
        "issue_date",
        "candidate.name",
        "employment.role",
        "employment.department",
        "employment.joining_date",
        "employment.employment_type",
        "employment.work_mode",
        "employment.location",
        "employment.working_days",
        "employment.working_hours",
        "employment.reporting_manager",
        "employment.compensation",
        "employment.responsibilities",
        "employment.equipment_and_access",
        "employment.expenses",
        "terms.leave_policy",
        "terms.notice_period",
        "terms.background_verification",
        "terms.governing_law",
    ]
    for field in required_fields:
        current: Any = offer
        for key in field.split("."):
            if not isinstance(current, dict) or key not in current:
                raise OfferValidationError(f"Missing required field: {field}")
            current = current[key]
        if current is None or current == "" or current == []:
            raise OfferValidationError(f"Required field is empty: {field}")
    # Reuse the profile and asset checks from the internship generator.
    validate(profile, {
        "reference_number": offer["reference_number"],
        "issue_date": offer["issue_date"],
        "candidate": offer["candidate"],
        "internship": {
            "role": offer["employment"]["role"],
            "department": offer["employment"]["department"],
            "start_date": offer["employment"]["joining_date"],
            "end_date": offer["employment"]["joining_date"],
            "work_mode": offer["employment"]["work_mode"],
            "location": offer["employment"]["location"],
            "working_days": offer["employment"]["working_days"],
            "working_hours": offer["employment"]["working_hours"],
            "reporting_manager": offer["employment"]["reporting_manager"],
            "compensation": offer["employment"]["compensation"],
            "responsibilities": offer["employment"]["responsibilities"],
            "equipment_and_access": offer["employment"]["equipment_and_access"],
            "expenses": offer["employment"]["expenses"],
            "completion_documents": "Not applicable",
        },
        "terms": {
            "notice_period": offer["terms"]["notice_period"],
            "leave_policy": offer["terms"]["leave_policy"],
            "background_verification": offer["terms"]["background_verification"],
            "governing_law": offer["terms"]["governing_law"],
        },
        "acceptance_deadline": offer["issue_date"],
    })


def add_employment_terms_table(document: Document, profile: dict[str, Any], offer: dict[str, Any]) -> None:
    employment = offer["employment"]
    manager = f"{employment['reporting_manager']} ({employment.get('reporting_manager_title', 'CEO')})"
    rows = [
        ("Role", employment["role"]),
        ("Department", employment["department"]),
        ("Employment type", employment["employment_type"]),
        ("Joining date", display_date(employment["joining_date"], "employment.joining_date")),
        ("Work arrangement", employment["work_mode"]),
        ("Working schedule", f"{employment['working_days']} | {employment['working_hours']}"),
        ("Reporting to", manager),
        ("Annual compensation", employment["compensation"]),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [3000, CONTENT_WIDTH_DXA - 3000])
    table.style = "Table Grid"
    for row_idx, (label, value) in enumerate(rows):
        left, right = table.rows[row_idx].cells
        from generate_offer_letter import set_cell_shading
        set_cell_shading(left, "F3F0FC")
        if row_idx % 2 == 1:
            set_cell_shading(right, "FBFAFE")
        for cell, text_value, bold, color in (
            (left, label, True, DARK_PURPLE),
            (right, value, False, INK),
        ):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(text_value)
            set_run_font(run, size=9, color=color, bold=bold)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(0)


def build_employment_document(profile: dict[str, Any], offer: dict[str, Any]) -> Document:
    document = Document()
    configure_styles(document)
    configure_page_and_furniture(document, profile, document_label="Employment Offer Letter")
    bullet_num_id = create_bullet_numbering(document)
    document.core_properties.title = f"Employment Offer Letter - {offer['candidate']['name']}"
    document.core_properties.subject = f"Offer for {offer['employment']['role']}"
    document.core_properties.author = profile["legal_name"]
    document.core_properties.company = profile["legal_name"]

    if offer.get("document_status"):
        status = document.add_paragraph()
        status.alignment = WD_ALIGN_PARAGRAPH.CENTER
        status.paragraph_format.space_after = Pt(4)
        run = status.add_run(offer["document_status"])
        set_run_font(run, size=7.5, color=WHITE, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    title_run = title.add_run("EMPLOYMENT OFFER LETTER")
    set_run_font(title_run, size=16, color=DARK_PURPLE, bold=True)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.paragraph_format.space_after = Pt(1)
    set_run_font(meta.add_run(display_date(offer["issue_date"], "issue_date")), size=9.5, color=INK)
    ref = document.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ref.paragraph_format.space_after = Pt(5)
    set_run_font(ref.add_run(f"Reference: {offer['reference_number']}"), size=9, color=MUTED)

    candidate = offer["candidate"]
    add_label_value_paragraph(document, "To", candidate["name"])
    subject = document.add_paragraph()
    subject.paragraph_format.space_after = Pt(4)
    set_run_font(subject.add_run(f"Subject: Offer of employment as {offer['employment']['role']}"), size=9.5, color=DARK_PURPLE, bold=True)
    salutation = document.add_paragraph(f"Dear {candidate['name']},")
    salutation.paragraph_format.space_after = Pt(4)
    intro = document.add_paragraph()
    intro.add_run("We are pleased to offer you full-time employment with ")
    company = intro.add_run(profile["legal_name"])
    company.bold = True
    intro.add_run(f" as {offer['employment']['role']}. This letter records the principal terms of your employment. Please read it carefully and retain a copy for your records.")

    document.add_paragraph("Key terms", style="Heading 1")
    add_employment_terms_table(document, profile, offer)

    add_term(document, "1. Appointment and employment", f"Your appointment as {offer['employment']['role']} will commence on {display_date(offer['employment']['joining_date'], 'employment.joining_date')}. This is a {offer['employment']['employment_type'].lower()} position. Your continued employment is subject to satisfactory performance, professional conduct, eligibility to work, and compliance with this letter and Nhancio policies.")
    document.add_paragraph("2. Role responsibilities and standards", style="Heading 2")
    for item in offer["employment"]["responsibilities"]:
        add_bullet(document, item, bullet_num_id)
    paragraph = document.add_paragraph("Responsibilities may reasonably evolve with business needs. You are expected to perform assigned work diligently, communicate blockers promptly, maintain accurate records, participate in reviews, and follow lawful instructions relevant to the role.")
    paragraph.paragraph_format.space_before = Pt(2)
    employment = offer["employment"]
    add_term(document, "3. Working arrangements, attendance, and leave", f"Your normal arrangement is {employment['work_mode'].lower()}. Work will be performed remotely on {employment['working_days']}, for {employment['working_hours']}. Reasonable schedule changes may be agreed with Didigam Nithin (CEO). {offer['terms']['leave_policy']} Persistent unapproved absence, material lateness, or failure to communicate availability may result in review or disciplinary action.")
    add_term(document, "4. Compensation, deductions, and expenses", f"Your annual compensation is {employment['compensation']}. Salary is subject to applicable deductions, taxes, payroll rules, and statutory requirements. {employment['expenses']}")
    add_term(document, "5. Reporting and performance", "You will report to Didigam Nithin (CEO). Nhancio may set goals, request status updates, review work quality, and provide feedback. Performance reviews, role progression, and any compensation changes are subject to company processes and business requirements.")
    add_term(document, "6. Confidentiality and information security", "During and after employment, you must keep confidential all non-public company, client, partner, employee, applicant, and user information learned through your work. Confidential information includes business plans, pricing, credentials, source code, models, prompts, datasets, designs, research, processes, and personal data. Use such information only for authorized work; do not copy it to personal systems, disclose it, publish it, or use it for an external portfolio without prior written approval. Immediately report suspected loss, unauthorized access, or disclosure. These duties survive the end of employment.")
    add_term(document, "7. Intellectual property and work product", "To the extent permitted by applicable law, work product created within the scope of assigned duties using Nhancio or client time, information, systems, or resources will belong to Nhancio or the relevant client, as directed. You agree to disclose such work product and reasonably assist with documentation needed to confirm ownership. Pre-existing intellectual property remains yours if disclosed in writing before it is incorporated into assigned work; its inclusion requires prior approval and an appropriate licence. No company or client material may be open-sourced, published, demonstrated, or included in a portfolio without written authorization.")
    add_term(document, "8. Conduct, conflicts, and policies", "You must act professionally, respectfully, lawfully, and without discrimination or harassment. Disclose any actual or potential conflict of interest, including overlapping work that could compromise confidentiality, availability, or ownership of work product. You must follow applicable Nhancio and client policies communicated to you, including responsible-AI, acceptable-use, security, privacy, and workplace-conduct requirements. You may not represent that you can bind Nhancio or make commitments on its behalf unless expressly authorized in writing.")
    add_term(document, "9. Systems, equipment, and return of property", employment["equipment_and_access"] + " On request or when employment ends, you must promptly return physical property, transfer current work, delete company or client information from personal devices and accounts where permitted, and complete offboarding steps.")
    add_term(document, "10. Verification and eligibility", offer["terms"]["background_verification"] + " You are responsible for maintaining any work authorization, identification, or other eligibility required for employment and for promptly disclosing any change that affects eligibility.")
    add_term(document, "11. Ending employment", f"Either you or Nhancio may end employment by giving {offer['terms']['notice_period']} written notice, unless a shorter period is mutually agreed. Nhancio may end employment immediately for serious misconduct, material breach of confidentiality or security, falsification, unlawful conduct, repeated non-performance, policy violations, loss of eligibility, or conduct that reasonably creates material risk to Nhancio, its clients, or users. On termination, salary and eligible dues will be handled according to the stated terms and applicable law for the period completed.")
    add_term(document, "12. Governing terms and changes", f"This letter and any policies expressly incorporated into it record the understanding for your employment and replace prior discussions about the same engagement. Any amendment must be in writing and approved by an authorized Nhancio representative. If a provision is unenforceable, the remaining provisions continue to apply to the extent permitted by law. This letter is governed by the laws applicable in {offer['terms']['governing_law']}, and the competent courts there will have jurisdiction, subject to any mandatory law that applies.")

    document.add_paragraph("Authorization", style="Heading 1")
    p = document.add_paragraph("We look forward to your contribution and professional growth with Nhancio.")
    p.paragraph_format.space_after = Pt(4)
    signatory = profile["signatory"]
    auth = document.add_paragraph()
    auth.paragraph_format.space_after = Pt(2)
    set_run_font(auth.add_run(f"For and on behalf of {profile['legal_name']}"), size=9.5, color=DARK_PURPLE, bold=True)
    sig_p = document.add_paragraph()
    sig_p.paragraph_format.space_after = Pt(2)
    shape = sig_p.add_run().add_picture(str(BASE_DIR / signatory["signature_image"]), width=Inches(1.25))
    add_picture_alt(shape, "Authorized signature", f"Signature of {signatory['name']}")
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    set_run_font(p.add_run(signatory["name"]), size=9.5, color=INK, bold=True)
    add_label_value_paragraph(document, "Title", signatory["title"])
    return document


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--offer", type=Path, default=BASE_DIR / "employment_sample_offer.json")
    parser.add_argument("--profile", type=Path, default=BASE_DIR / "company_profile.json")
    parser.add_argument("--output", type=Path)
    parser.add_argument("--name")
    parser.add_argument("--role")
    parser.add_argument("--joining-date", dest="joining_date")
    parser.add_argument("--salary-lpa", dest="salary_lpa", type=float)
    parser.add_argument("--leave-policy", dest="leave_policy")
    parser.add_argument("--issue-date", dest="issue_date")
    parser.add_argument("--reference")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    try:
        profile = load_json(args.profile.resolve())
        offer = load_json(args.offer.resolve())
        if args.name:
            offer["candidate"]["name"] = args.name
        if args.role:
            offer["employment"]["role"] = args.role
        if args.joining_date:
            offer["employment"]["joining_date"] = args.joining_date
        if args.salary_lpa is not None:
            annual = int(round(args.salary_lpa * 100000))
            monthly = annual / 12
            offer["employment"]["compensation"] = f"INR {annual:,.0f} per annum ({args.salary_lpa:g} LPA), equivalent to approximately INR {monthly:,.0f} per month before applicable deductions"
        if args.leave_policy:
            offer["terms"]["leave_policy"] = args.leave_policy
        if args.issue_date:
            offer["issue_date"] = args.issue_date
        if args.reference:
            offer["reference_number"] = args.reference
        validate_employment(profile, offer)
        output = (args.output or BASE_DIR / "generated" / f"{safe_filename(offer['candidate']['name'])}_Senior_Developer_Offer_Letter.docx").resolve()
        output.parent.mkdir(parents=True, exist_ok=True)
        build_employment_document(profile, offer).save(output)
        print(output)
        return 0
    except OfferValidationError as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
