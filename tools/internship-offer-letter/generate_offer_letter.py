#!/usr/bin/env python3
"""Generate a branded Nhancio internship offer letter from JSON data."""

from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import date, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
DEFAULT_PROFILE = BASE_DIR / "company_profile.json"
DEFAULT_OUTPUT_DIR = BASE_DIR / "generated"

PURPLE = RGBColor(91, 45, 203)
DARK_PURPLE = RGBColor(50, 30, 96)
CYAN = RGBColor(0, 180, 216)
INK = RGBColor(31, 31, 36)
MUTED = RGBColor(96, 99, 110)
LIGHT_FILL = "F3F0FC"
LIGHT_GRAY = "F4F5F7"
WHITE = RGBColor(255, 255, 255)

# Resolved `standard_business_brief` preset with a named Nhancio brand override.
# A4, compact branded layout tuned to keep the comprehensive letter to two pages.
PAGE_WIDTH_DXA = 11906
PAGE_HEIGHT_DXA = 16838
MARGIN_DXA = 792
CONTENT_WIDTH_DXA = PAGE_WIDTH_DXA - (2 * MARGIN_DXA)
TABLE_INDENT_DXA = 0
CELL_MARGINS_DXA = {"top": 55, "bottom": 55, "start": 100, "end": 100}


class OfferValidationError(ValueError):
    """Raised when an offer file is incomplete or internally inconsistent."""


def load_json(path: Path) -> dict[str, Any]:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise OfferValidationError(f"File not found: {path}") from exc
    except json.JSONDecodeError as exc:
        raise OfferValidationError(f"Invalid JSON in {path}: {exc}") from exc


def required(data: dict[str, Any], dotted_path: str) -> Any:
    current: Any = data
    for key in dotted_path.split("."):
        if not isinstance(current, dict) or key not in current:
            raise OfferValidationError(f"Missing required field: {dotted_path}")
        current = current[key]
    if current is None or current == "" or current == []:
        raise OfferValidationError(f"Required field is empty: {dotted_path}")
    return current


def parse_iso_date(value: str, field_name: str) -> date:
    try:
        return datetime.strptime(value, "%Y-%m-%d").date()
    except (TypeError, ValueError) as exc:
        raise OfferValidationError(
            f"{field_name} must use YYYY-MM-DD format; received {value!r}"
        ) from exc


def display_date(value: str, field_name: str) -> str:
    parsed = parse_iso_date(value, field_name)
    return f"{parsed.day} {parsed.strftime('%B %Y')}"


def safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    return cleaned.strip("._") or "candidate"


def validate(profile: dict[str, Any], offer: dict[str, Any]) -> None:
    profile_fields = [
        "display_name",
        "legal_name",
        "website",
        "address_lines",
        "logo_image",
        "signatory.name",
        "signatory.title",
        "signatory.signature_image",
        "internship_supervisor.name",
        "internship_supervisor.title",
    ]
    offer_fields = [
        "reference_number",
        "issue_date",
        "candidate.name",
        "internship.role",
        "internship.department",
        "internship.start_date",
        "internship.end_date",
        "internship.work_mode",
        "internship.location",
        "internship.working_days",
        "internship.working_hours",
        "internship.reporting_manager",
        "internship.compensation",
        "internship.responsibilities",
        "internship.equipment_and_access",
        "internship.expenses",
        "internship.completion_documents",
        "terms.notice_period",
        "terms.leave_policy",
        "terms.background_verification",
        "terms.governing_law",
        "acceptance_deadline",
    ]
    for field in profile_fields:
        required(profile, field)
    for field in offer_fields:
        required(offer, field)

    issue = parse_iso_date(offer["issue_date"], "issue_date")
    start = parse_iso_date(offer["internship"]["start_date"], "internship.start_date")
    end = parse_iso_date(offer["internship"]["end_date"], "internship.end_date")
    accept_by = parse_iso_date(offer["acceptance_deadline"], "acceptance_deadline")
    if end < start:
        raise OfferValidationError("internship.end_date cannot be before internship.start_date")
    if accept_by < issue:
        raise OfferValidationError("acceptance_deadline cannot be before issue_date")

    responsibilities = offer["internship"]["responsibilities"]
    if not isinstance(responsibilities, list) or not all(
        isinstance(item, str) and item.strip() for item in responsibilities
    ):
        raise OfferValidationError("internship.responsibilities must be a non-empty list of text items")

    for key in ("logo_image",):
        asset_path = (BASE_DIR / profile[key]).resolve()
        if not asset_path.is_file():
            raise OfferValidationError(f"Missing asset: {asset_path}")
    signature_path = (BASE_DIR / profile["signatory"]["signature_image"]).resolve()
    if not signature_path.is_file():
        raise OfferValidationError(f"Missing asset: {signature_path}")


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0

    heading_tokens = {
        "Heading 1": (13, PURPLE, 7, 3),
        "Heading 2": (10.5, PURPLE, 5, 2),
        "Heading 3": (10, DARK_PURPLE, 4, 2),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGINS_DXA.items():
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges: dict[str, Any]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, values in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        for key in ("val", "sz", "space", "color"):
            if key in values:
                edge.set(qn(f"w:{key}"), str(values[key]))


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA} DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_picture_alt(inline_shape, title: str, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_page_field(paragraph) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "6B6E78")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "17")
    r_pr.extend([color, size])
    text = OxmlElement("w:t")
    text.text = "1"
    run.extend([r_pr, text])
    field.append(run)
    paragraph._p.append(field)


def configure_page_and_furniture(
    document: Document,
    profile: dict[str, Any],
    document_label: str = "Internship Offer Letter",
) -> None:
    section = document.sections[0]
    section.page_width = Inches(PAGE_WIDTH_DXA / 1440)
    section.page_height = Inches(PAGE_HEIGHT_DXA / 1440)
    section.top_margin = Inches(MARGIN_DXA / 1440)
    section.bottom_margin = Inches(MARGIN_DXA / 1440)
    section.left_margin = Inches(MARGIN_DXA / 1440)
    section.right_margin = Inches(MARGIN_DXA / 1440)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)

    header = section.header
    p_logo = header.paragraphs[0]
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(0)
    p_logo.paragraph_format.line_spacing = 0.1
    logo_run = p_logo.add_run(" ")
    set_run_font(logo_run, size=1, color=WHITE)
    header_table = header.add_table(rows=1, cols=2, width=Inches(CONTENT_WIDTH_DXA / 1440))
    set_table_geometry(header_table, [1900, CONTENT_WIDTH_DXA - 1900], indent_dxa=0)
    logo_cell, brand_cell = header_table.rows[0].cells
    for cell in header_table.rows[0].cells:
        set_cell_border(
            cell,
            top={"val": "nil"}, bottom={"val": "nil"},
            start={"val": "nil"}, end={"val": "nil"},
        )
    logo_p = logo_cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_p.paragraph_format.space_after = Pt(0)
    logo_run = logo_p.add_run()
    shape = logo_run.add_picture(str(BASE_DIR / profile["logo_image"]), width=Inches(0.46))
    add_picture_alt(shape, "Nhancio logo", "Nhancio brand logo")
    brand_p = brand_cell.paragraphs[0]
    brand_p.paragraph_format.space_after = Pt(0)
    brand = brand_p.add_run(profile["display_name"].upper())
    set_run_font(brand, size=11.5, color=PURPLE, bold=True)
    brand_p.add_run("  ✦  ")
    site = brand_p.add_run(profile["website"].replace("https://", ""))
    set_run_font(site, size=9, color=MUTED)
    details = brand_cell.add_paragraph()
    details.paragraph_format.space_before = Pt(0)
    details.paragraph_format.space_after = Pt(0)
    detail_run = details.add_run("✦  Incubated at T-HUB, Hyd   ✦  AI Deployment Services")
    set_run_font(detail_run, size=8, color=MUTED)

    footer = section.footer
    p_footer = footer.paragraphs[0]
    p_footer.paragraph_format.space_before = Pt(0)
    p_footer.paragraph_format.space_after = Pt(0)
    p_footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.1))
    left = p_footer.add_run(f"{profile['display_name']} | Confidential | {document_label}")
    set_run_font(left, size=8.5, color=MUTED)
    page_label = p_footer.add_run("\tPage ")
    set_run_font(page_label, size=8.5, color=MUTED)
    add_page_field(p_footer)


def next_numbering_id(numbering, tag: str) -> int:
    values = []
    attr = qn("w:abstractNumId") if tag == "w:abstractNum" else qn("w:numId")
    for element in numbering.findall(qn(tag)):
        try:
            values.append(int(element.get(attr)))
        except (TypeError, ValueError):
            continue
    return max(values, default=0) + 1


def create_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "w:abstractNum")
    num_id = next_numbering_id(numbering, "w:num")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.extend([start, num_fmt, lvl_text, lvl_jc])

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(document: Document, text: str, num_id: int) -> None:
    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5, color=INK)


def add_label_value_paragraph(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=9.5, color=INK, bold=True)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, size=9.5, color=INK)


def add_key_terms_table(document: Document, profile: dict[str, Any], offer: dict[str, Any]) -> None:
    internship = offer["internship"]
    supervisor = profile["internship_supervisor"]
    manager = f"{supervisor['name']} ({supervisor['title']})"
    arrangement = internship["work_mode"]
    if internship.get("location") and internship["location"].lower() != internship["work_mode"].lower():
        arrangement = f"{arrangement} | {internship['location']}"
    rows = [
        ("Role", internship["role"]),
        ("Department", internship["department"]),
        (
            "Internship period",
            f"{display_date(internship['start_date'], 'internship.start_date')} to "
            f"{display_date(internship['end_date'], 'internship.end_date')}",
        ),
            ("Work arrangement", arrangement),
        ("Working schedule", f"{internship['working_days']} | {internship['working_hours']}"),
        ("Reporting to", manager),
        ("Compensation", internship["compensation"]),
        (
            "Accept by",
            display_date(offer["acceptance_deadline"], "acceptance_deadline"),
        ),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [3000, CONTENT_WIDTH_DXA - 3000])
    table.style = "Table Grid"
    for row_idx, (label, value) in enumerate(rows):
        left, right = table.rows[row_idx].cells
        set_cell_shading(left, LIGHT_FILL)
        if row_idx % 2 == 1:
            set_cell_shading(right, "FBFAFE")
        for cell, text_value, bold, color in (
            (left, label, True, DARK_PURPLE),
            (right, value, False, INK),
        ):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run(text_value)
            set_run_font(run, size=9, color=color, bold=bold)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(0)


def add_term(document: Document, heading: str, body: str) -> None:
    document.add_paragraph(heading, style="Heading 2")
    paragraph = document.add_paragraph(body)
    paragraph.paragraph_format.keep_together = False


def add_signature_lines(document: Document) -> None:
    table = document.add_table(rows=2, cols=2)
    set_table_geometry(table, [CONTENT_WIDTH_DXA - 3120, 3120])
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell in table.rows[0].cells:
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(" ")
        set_cell_border(cell, bottom={"val": "single", "sz": 8, "space": 0, "color": "6B6E78"})
    labels = ("Candidate signature", "Date")
    for idx, label in enumerate(labels):
        cell = table.rows[1].cells[idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run_font(run, size=9, color=MUTED)
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"})


def build_document(profile: dict[str, Any], offer: dict[str, Any]) -> Document:
    document = Document()
    configure_styles(document)
    configure_page_and_furniture(document, profile)
    bullet_num_id = create_bullet_numbering(document)

    document.core_properties.title = f"Internship Offer Letter - {offer['candidate']['name']}"
    document.core_properties.subject = f"Offer for {offer['internship']['role']}"
    document.core_properties.author = profile["legal_name"]
    document.core_properties.company = profile["legal_name"]
    document.core_properties.comments = "Generated from the Nhancio internship offer-letter template."

    if offer.get("document_status"):
        status = document.add_paragraph()
        status.alignment = WD_ALIGN_PARAGRAPH.CENTER
        status.paragraph_format.space_after = Pt(4)
        run = status.add_run(offer["document_status"])
        set_run_font(run, size=7.5, color=WHITE, bold=True)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "6B2FD0")
        status._p.get_or_add_pPr().append(shading)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(5)
    title_run = title.add_run("INTERNSHIP OFFER LETTER")
    set_run_font(title_run, size=16, color=DARK_PURPLE, bold=True)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.paragraph_format.space_after = Pt(1)
    r = meta.add_run(display_date(offer["issue_date"], "issue_date"))
    set_run_font(r, size=10.5, color=INK)
    ref = document.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ref.paragraph_format.space_after = Pt(5)
    r = ref.add_run(f"Reference: {offer['reference_number']}")
    set_run_font(r, size=9.5, color=MUTED)

    candidate = offer["candidate"]
    add_label_value_paragraph(document, "To", candidate["name"])
    if candidate.get("address_lines"):
        for line in candidate["address_lines"]:
            p = document.add_paragraph(line)
            p.paragraph_format.space_after = Pt(1)
    contact_parts = [value for value in (candidate.get("email"), candidate.get("phone")) if value]
    if contact_parts:
        p = document.add_paragraph(" | ".join(contact_parts))
        p.paragraph_format.space_after = Pt(4)

    subject = document.add_paragraph()
    subject.paragraph_format.space_after = Pt(4)
    r = subject.add_run(f"Subject: Offer of internship as {offer['internship']['role']}")
    set_run_font(r, size=9.5, color=DARK_PURPLE, bold=True)

    salutation = document.add_paragraph(f"Dear {candidate['name']},")
    salutation.paragraph_format.space_after = Pt(4)
    intro = document.add_paragraph()
    intro.add_run("We are pleased to offer you an internship with ")
    company_run = intro.add_run(profile["legal_name"])
    company_run.bold = True
    intro.add_run(
        f" as {offer['internship']['role']}. This letter records the principal terms of the internship. "
        "Please read it carefully and retain a copy for your records."
    )

    document.add_paragraph("Key terms", style="Heading 1")
    add_key_terms_table(document, profile, offer)

    add_term(
        document,
        "1. Nature and purpose of the internship",
        "The internship is a structured, time-bound learning engagement intended to provide practical exposure, supervised project experience, and professional development. It does not guarantee future employment, conversion to a permanent role, or continued engagement after the stated end date. Any extension or employment offer must be made separately in writing by an authorized representative of Nhancio.",
    )

    document.add_paragraph("2. Responsibilities and expected standards", style="Heading 2")
    for item in offer["internship"]["responsibilities"]:
        add_bullet(document, item, bullet_num_id)
    paragraph = document.add_paragraph(
        "Assignments may reasonably evolve during the internship. You are expected to perform agreed work diligently, communicate blockers promptly, maintain accurate records, participate in reviews, and follow lawful instructions relevant to the role."
    )
    paragraph.paragraph_format.space_before = Pt(2)

    add_term(
        document,
        "3. Working arrangements, attendance, and leave",
        f"Your normal arrangement is {offer['internship']['work_mode'].lower()}. Work will be performed remotely on "
        f"{offer['internship']['working_days']}, for {offer['internship']['working_hours']}. Reasonable schedule changes may be agreed with Didigam Nithin (CEO). "
        f"{offer['terms']['leave_policy']} Persistent unapproved absence, material lateness, or failure to communicate availability may result in review or termination of the internship.",
    )

    add_term(
        document,
        "4. Compensation, taxes, and expenses",
        f"The compensation for this internship is: {offer['internship']['compensation']}. You are responsible for providing accurate payment and tax information and for personal tax obligations that apply to you. "
        f"{offer['internship']['expenses']}",
    )

    add_term(
        document,
        "5. Supervision and performance",
        "You will report to Didigam Nithin (CEO). Nhancio may set milestones, request status updates, review work quality, and provide feedback. Continued participation and any completion documentation depend on satisfactory performance, attendance, professional conduct, and fulfillment of handover obligations.",
    )

    add_term(
        document,
        "6. Confidentiality and information security",
        "During and after the internship, you must keep confidential all non-public company, client, partner, employee, applicant, and user information learned through the engagement. Confidential information includes business plans, pricing, credentials, source code, models, prompts, datasets, designs, research, processes, and personal data. Use such information only for authorized internship work; do not copy it to personal systems, disclose it, publish it, or use it for an external portfolio without prior written approval. Immediately report suspected loss, unauthorized access, or disclosure and follow all access-control and data-handling instructions. These duties survive the end of the internship.",
    )

    add_term(
        document,
        "7. Intellectual property and work product",
        "To the extent permitted by applicable law, work product created specifically within the scope of assigned internship duties using Nhancio or client time, information, systems, or resources will belong to Nhancio or the relevant client, as directed. You agree to disclose such work product and reasonably assist with documentation needed to confirm ownership. Pre-existing intellectual property remains yours if disclosed in writing before it is incorporated into assigned work; its inclusion requires prior approval and an appropriate licence. No company or client material may be open-sourced, published, demonstrated, or included in a portfolio without written authorization.",
    )

    add_term(
        document,
        "8. Conduct, conflicts, and policies",
        "You must act professionally, respectfully, lawfully, and without discrimination or harassment. Disclose any actual or potential conflict of interest, including overlapping work that could compromise confidentiality, availability, or ownership of work product. You must follow applicable Nhancio and client policies communicated to you, including responsible-AI, acceptable-use, security, privacy, and workplace-conduct requirements. You may not represent that you can bind Nhancio or make commitments on its behalf unless expressly authorized in writing.",
    )

    add_term(
        document,
        "9. Systems, equipment, and return of property",
        f"{offer['internship']['equipment_and_access']} On request or at the end of the internship, you must promptly return physical property, transfer current work, delete company or client information from personal devices and accounts where permitted, and confirm completion of the offboarding steps.",
    )

    add_term(
        document,
        "10. Verification and eligibility",
        offer["terms"]["background_verification"]
        + " You are responsible for maintaining any academic permission, work authorization, identification, or other eligibility required for the internship and for promptly disclosing any change that affects eligibility.",
    )

    add_term(
        document,
        "11. Ending the internship",
        f"Either you or Nhancio may end the internship by giving {offer['terms']['notice_period']} written notice, unless a shorter period is mutually agreed. Nhancio may end the internship immediately for serious misconduct, material breach of confidentiality or security, falsification, unlawful conduct, repeated non-performance, policy violations, loss of eligibility, or conduct that reasonably creates material risk to Nhancio, its clients, or users. On termination, compensation, if any, will be handled according to the stated terms and applicable law for the eligible period completed.",
    )

    add_term(
        document,
        "12. Completion documentation",
        offer["internship"]["completion_documents"],
    )

    add_term(
        document,
        "13. Governing terms and changes",
        f"This letter and any policies expressly incorporated into it record the understanding for the internship and replace prior discussions about the same engagement. Any amendment must be in writing and approved by an authorized Nhancio representative. If a provision is unenforceable, the remaining provisions continue to apply to the extent permitted by law. This letter is governed by the laws applicable in {offer['terms']['governing_law']}, and the competent courts there will have jurisdiction, subject to any mandatory law that applies.",
    )

    acceptance_heading = document.add_paragraph("Authorization", style="Heading 1")
    acceptance_heading.paragraph_format.keep_with_next = True
    p = document.add_paragraph("We look forward to the learning, contribution, and professional growth this internship can provide.")
    p.paragraph_format.space_after = Pt(4)

    signatory = profile["signatory"]
    auth = document.add_paragraph()
    auth.paragraph_format.space_after = Pt(2)
    r = auth.add_run(f"For and on behalf of {profile['legal_name']}")
    set_run_font(r, size=11, color=DARK_PURPLE, bold=True)

    sig_p = document.add_paragraph()
    sig_p.paragraph_format.space_before = Pt(0)
    sig_p.paragraph_format.space_after = Pt(2)
    shape = sig_p.add_run().add_picture(
        str(BASE_DIR / signatory["signature_image"]), width=Inches(1.25)
    )
    add_picture_alt(shape, "Authorized signature", f"Signature of {signatory['name']}")

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(signatory["name"])
    set_run_font(r, size=11, color=INK, bold=True)
    add_label_value_paragraph(document, "Title", signatory["title"])
    if signatory.get("phone"):
        add_label_value_paragraph(document, "Phone", signatory["phone"])

    return document


def apply_overrides(offer: dict[str, Any], args: argparse.Namespace) -> None:
    mapping = {
        "name": ("candidate", "name"),
        "role": ("internship", "role"),
        "start_date": ("internship", "start_date"),
        "end_date": ("internship", "end_date"),
        "issue_date": ("issue_date",),
        "accept_by": ("acceptance_deadline",),
        "reference": ("reference_number",),
    }
    for arg_name, path in mapping.items():
        value = getattr(args, arg_name)
        if value is None:
            continue
        current = offer
        for key in path[:-1]:
            current = current[key]
        current[path[-1]] = value


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--offer", required=True, type=Path, help="Completed offer JSON file")
    parser.add_argument("--profile", type=Path, default=DEFAULT_PROFILE, help="Company/signatory JSON")
    parser.add_argument("--output", type=Path, help="Destination .docx path")
    parser.add_argument("--name", help="Override candidate name")
    parser.add_argument("--role", help="Override internship role")
    parser.add_argument("--start-date", dest="start_date", help="Override start date (YYYY-MM-DD)")
    parser.add_argument("--end-date", dest="end_date", help="Override end date (YYYY-MM-DD)")
    parser.add_argument("--issue-date", dest="issue_date", help="Override issue date (YYYY-MM-DD)")
    parser.add_argument("--accept-by", dest="accept_by", help="Override acceptance deadline (YYYY-MM-DD)")
    parser.add_argument("--reference", help="Override reference number")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    try:
        profile = load_json(args.profile.resolve())
        offer = load_json(args.offer.resolve())
        apply_overrides(offer, args)
        validate(profile, offer)
        output = args.output
        if output is None:
            output = DEFAULT_OUTPUT_DIR / (
                f"{safe_filename(offer['candidate']['name'])}_Internship_Offer_Letter.docx"
            )
        output = output.resolve()
        if output.suffix.lower() != ".docx":
            raise OfferValidationError("Output filename must end in .docx")
        output.parent.mkdir(parents=True, exist_ok=True)
        document = build_document(profile, offer)
        document.save(output)
        print(output)
        return 0
    except OfferValidationError as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
